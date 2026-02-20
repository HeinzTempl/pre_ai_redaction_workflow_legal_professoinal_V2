from docx import Document
import re
import json
import os
from llm_api import redact_text_api

# ==================== LERNEBENE ====================
# Persistente Korrekturliste: Begriffe die immer/nie geschwärzt werden sollen

LEARNED_ENTITIES_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "learned_entities.json")

_learned_data = {
    "never_redact": [],       # Begriffe die NIE geschwärzt werden (False Positives)
    "always_redact": {        # Begriffe die IMMER geschwärzt werden (Missed Entities)
        "PER": [],
        "ORG": [],
        "LOC": [],
    }
}


def load_learned_entities():
    """Lädt die gelernte Entity-Liste aus der JSON-Datei."""
    global _learned_data
    if os.path.exists(LEARNED_ENTITIES_FILE):
        try:
            with open(LEARNED_ENTITIES_FILE, "r", encoding="utf-8") as f:
                _learned_data = json.load(f)
            print(f"  Gelernte Entities geladen: {LEARNED_ENTITIES_FILE}")
        except Exception as e:
            print(f"  Warnung: Konnte gelernte Entities nicht laden: {e}")


def save_learned_entities():
    """Speichert die gelernte Entity-Liste in die JSON-Datei."""
    try:
        with open(LEARNED_ENTITIES_FILE, "w", encoding="utf-8") as f:
            json.dump(_learned_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"  Warnung: Konnte gelernte Entities nicht speichern: {e}")


def add_never_redact(text):
    """Fügt einen Begriff zur 'nie schwärzen'-Liste hinzu."""
    text = text.strip()
    if text and text not in _learned_data["never_redact"]:
        _learned_data["never_redact"].append(text)
        save_learned_entities()


def remove_never_redact(text):
    """Entfernt einen Begriff von der 'nie schwärzen'-Liste."""
    text = text.strip()
    if text in _learned_data["never_redact"]:
        _learned_data["never_redact"].remove(text)
        save_learned_entities()


def add_always_redact(text, label="PER"):
    """Fügt einen Begriff zur 'immer schwärzen'-Liste hinzu."""
    text = text.strip()
    if label not in _learned_data["always_redact"]:
        _learned_data["always_redact"][label] = []
    if text and text not in _learned_data["always_redact"][label]:
        _learned_data["always_redact"][label].append(text)
        save_learned_entities()


def remove_always_redact(text, label="PER"):
    """Entfernt einen Begriff von der 'immer schwärzen'-Liste."""
    text = text.strip()
    if label in _learned_data["always_redact"] and text in _learned_data["always_redact"][label]:
        _learned_data["always_redact"][label].remove(text)
        save_learned_entities()


def is_learned_never_redact(text):
    """Prüft ob ein Begriff auf der 'nie schwärzen'-Liste steht."""
    return text.strip() in _learned_data["never_redact"]


def get_learned_always_redact():
    """Gibt alle 'immer schwärzen'-Begriffe zurück."""
    return _learned_data.get("always_redact", {})


def get_learned_data():
    """Gibt die gesamte gelernte Datenliste zurück (für UI)."""
    return _learned_data


# Beim Import laden
load_learned_entities()

# ==================== NER-ENGINE LADEN ====================
# Unterstützte Engines: "flair" (Standard, genauer) und "spacy" (schneller, optional)

_nlp_engine = None
_nlp_engine_name = None
_flair_tagger_legal = None
_flair_tagger_large = None
_spacy_nlp = None


def load_flair_models():
    """Lädt die Flair NER-Modelle (legal + large)."""
    global _flair_tagger_legal, _flair_tagger_large, _nlp_engine, _nlp_engine_name
    from flair.models import SequenceTagger

    try:
        _flair_tagger_legal = SequenceTagger.load("flair/ner-german-legal")
        print("  Flair-Modell geladen: ner-german-legal")
    except Exception as e:
        print(f"  Warnung: ner-german-legal konnte nicht geladen werden: {e}")
        _flair_tagger_legal = None

    try:
        _flair_tagger_large = SequenceTagger.load("flair/ner-german-large")
        print("  Flair-Modell geladen: ner-german-large")
    except Exception as e:
        print(f"  Warnung: ner-german-large konnte nicht geladen werden: {e}")
        _flair_tagger_large = None

    if _flair_tagger_legal or _flair_tagger_large:
        _nlp_engine = "flair"
        _nlp_engine_name = "Flair (legal + large)"
    else:
        print("  Keine Flair-Modelle verfügbar, falle auf spaCy zurück.")
        load_spacy_model()


def load_spacy_model():
    """Lädt das spaCy-Modell als Fallback oder schnelle Alternative."""
    global _spacy_nlp, _nlp_engine, _nlp_engine_name
    import spacy
    try:
        _spacy_nlp = spacy.load("de_core_news_lg")
        print("  spaCy-Modell geladen: de_core_news_lg")
    except OSError:
        try:
            _spacy_nlp = spacy.load("de_core_news_md")
            print("  spaCy-Modell geladen: de_core_news_md")
        except OSError:
            _spacy_nlp = spacy.load("de_core_news_sm")
            print("  spaCy-Modell geladen: de_core_news_sm (Qualität eingeschränkt)")
    _nlp_engine = "spacy"
    _nlp_engine_name = "spaCy"


def set_ner_engine(engine="flair"):
    """
    Wählt die NER-Engine.
    engine: "flair" (Standard, genauer) oder "spacy" (schneller)
    """
    global _nlp_engine
    if engine == "flair":
        if _flair_tagger_legal is None and _flair_tagger_large is None:
            load_flair_models()
        else:
            _nlp_engine = "flair"
    elif engine == "spacy":
        if _spacy_nlp is None:
            load_spacy_model()
        else:
            _nlp_engine = "spacy"
    print(f"  NER-Engine: {_nlp_engine}")


def get_engine_name():
    """Gibt den Namen der aktiven NER-Engine zurück."""
    return _nlp_engine_name or "nicht geladen"


# Standard: Flair laden
print("\nLade NER-Modelle...")
try:
    load_flair_models()
except ImportError:
    print("  Flair nicht installiert. Verwende spaCy.")
    load_spacy_model()
except Exception as e:
    print(f"  Fehler beim Laden von Flair: {e}. Verwende spaCy.")
    load_spacy_model()


# ==================== SENSITIVITÄTSSTUFEN ====================

SENSITIVITY_THRESHOLDS = {
    "konservativ": 0.90,
    "standard":    0.80,
    "aggressiv":   0.60,
}


# ==================== WHITELIST ====================

WHITELIST_ORGS = {
    "Amtsgericht", "Landesgericht", "Oberlandesgericht", "Bundesgerichtshof",
    "Bundesverfassungsgericht", "Bundesverwaltungsgericht", "Bundesfinanzhof",
    "Bundesarbeitsgericht", "Bundessozialgericht", "Verwaltungsgericht",
    "Verwaltungsgerichtshof", "Finanzgericht", "Sozialgericht", "Arbeitsgericht",
    "Landesarbeitsgericht", "Landessozialgericht", "Oberverwaltungsgericht",
    "Bezirksgericht", "Handelsgericht", "Oberster Gerichtshof", "Verfassungsgerichtshof",
    "Bundesgericht", "Kantonsgericht", "Obergericht",
    "Europäischer Gerichtshof", "EuGH", "EGMR",
    "Europäischer Gerichtshof für Menschenrechte",
    "Finanzamt", "Grundbuchamt", "Handelsregister", "Firmenbuch",
    "Standesamt", "Bezirkshauptmannschaft", "Magistrat",
    "Bundesministerium", "Landesregierung", "Bezirksregierung",
    "Staatsanwaltschaft", "Generalstaatsanwaltschaft",
    "Datenschutzbehörde", "Bundespolizei", "Polizei",
    "Bundesnetzagentur", "Kartellamt", "Bundeskartellamt",
    "Europäische Union", "EU", "Europäische Kommission",
    "Europäisches Parlament", "Europarat",
    "Vereinte Nationen", "UN", "NATO",
    "Bundesrepublik Deutschland", "Republik Österreich",
    "Schweizerische Eidgenossenschaft",
}

WHITELIST_LOCS = {
    "Deutschland", "Österreich", "Schweiz", "Liechtenstein", "Luxemburg",
    "Frankreich", "Italien", "Spanien", "Niederlande", "Belgien",
    "Großbritannien", "England", "Schottland", "Irland",
    "USA", "Vereinigte Staaten", "China", "Japan", "Russland",
    "Europa", "Asien", "Afrika", "Nordamerika", "Südamerika",
    "Bayern", "Baden-Württemberg", "Hessen", "Nordrhein-Westfalen",
    "Niedersachsen", "Sachsen", "Thüringen", "Brandenburg",
    "Mecklenburg-Vorpommern", "Sachsen-Anhalt", "Schleswig-Holstein",
    "Rheinland-Pfalz", "Saarland", "Berlin", "Hamburg", "Bremen",
    "Wien", "Niederösterreich", "Oberösterreich", "Steiermark",
    "Tirol", "Kärnten", "Salzburg", "Vorarlberg", "Burgenland",
    "Zürich", "Bern", "Luzern", "Basel", "Genf", "Lausanne",
}

WHITELIST_MISC = {
    "BGB", "ZPO", "StGB", "StPO", "HGB", "GmbHG", "AktG", "InsO",
    "ABGB", "UGB", "DSGVO", "BDSG", "DSG", "GDPR",
    "GmbH", "AG", "KG", "OG", "OHG", "e.V.", "e.G.",
    "GmbH & Co. KG", "UG", "SE",
}

COMMON_FALSE_POSITIVES = {
    "Kläger", "Beklagte", "Beklagter", "Antragsteller", "Antragsgegner",
    "Beschuldigte", "Beschuldigter", "Angeklagte", "Angeklagter",
    "Klägerseite", "Beklagtenseite", "Nebenintervenientin", "Nebenintervenient",
    "Berufungswerber", "Berufungswerberin", "Revisionswerber",
    "Beschwerdeführer", "Beschwerdeführerin",
    "Erblasser", "Erblasserin", "Erben", "Erbin",
    "Mieter", "Vermieter", "Käufer", "Verkäufer",
    "Arbeitgeber", "Arbeitnehmer", "Dienstgeber", "Dienstnehmer",
    "Gläubiger", "Schuldner", "Bürge",
    "Richter", "Richterin", "Staatsanwalt", "Staatsanwältin",
    "Rechtsanwalt", "Rechtsanwältin", "Notar", "Notarin",
    "Vorsitzende", "Vorsitzender", "Beisitzer",
    "Zeuge", "Zeugin", "Sachverständige", "Sachverständiger",
    "Bundesrepublik", "Republik",
    "Partei", "Parteien", "Vertragspartei",
    "Absatz", "Ziffer", "Satz", "Nummer",
}


def is_whitelisted(entity_text, entity_label):
    """Prüft ob eine Entity auf der Whitelist steht."""
    text_clean = entity_text.strip()
    if text_clean in WHITELIST_MISC:
        return True
    if entity_label == "ORG":
        if text_clean in WHITELIST_ORGS:
            return True
        for term in WHITELIST_ORGS:
            if text_clean.startswith(term) or term in text_clean:
                return True
    if entity_label == "LOC":
        if text_clean in WHITELIST_LOCS:
            return True
    return False


def _is_grundbuch_fraction(text):
    """Erkennt Grundbuch-Anteile wie 128/542, 1/3, 25/100 etc."""
    return bool(re.match(r'^\d{1,6}/\d{1,6}$', text.strip()))


def _should_skip_entity(ent_text, ent_label):
    """Zusätzliche Heuristiken um False Positives zu vermeiden."""
    text = ent_text.strip()
    if len(text) <= 1:
        return True
    if text.replace(" ", "").isdigit():
        return True
    if _is_grundbuch_fraction(text):
        return True
    if text in COMMON_FALSE_POSITIVES:
        return True
    return False


# ==================== ENTITY MAPPER ====================

class EntityMapper:
    def __init__(self, sensitivity="standard"):
        self.person_mapping = {}
        self.org_mapping = {}
        self.loc_mapping = {}
        self.person_counter = 0
        self.org_counter = 0
        self.loc_counter = 0
        self.sensitivity = sensitivity
        self.confidence_threshold = SENSITIVITY_THRESHOLDS.get(sensitivity, 0.80)
        self.skipped_whitelist = []
        self.skipped_low_confidence = []
        self.skipped_org_juristic = []  # Juristische Personen (bei konservativ übersprungen)

    def get_placeholder(self, entity_text, entity_label):
        entity_text_clean = entity_text.strip()
        if not entity_text_clean:
            return None
        if entity_label == "PER":
            if entity_text_clean not in self.person_mapping:
                self.person_mapping[entity_text_clean] = f"Person {chr(65 + self.person_counter)}"
                self.person_counter += 1
            return self.person_mapping[entity_text_clean]
        elif entity_label == "ORG":
            if entity_text_clean not in self.org_mapping:
                self.org_mapping[entity_text_clean] = f"Firma {chr(65 + self.org_counter)}"
                self.org_counter += 1
            return self.org_mapping[entity_text_clean]
        elif entity_label == "LOC":
            if entity_text_clean not in self.loc_mapping:
                self.loc_mapping[entity_text_clean] = f"Ort {chr(65 + self.loc_counter)}"
                self.loc_counter += 1
            return self.loc_mapping[entity_text_clean]
        return None


# ==================== REGEX-MUSTER ====================

def get_regex_patterns(sensitivity="standard"):
    patterns = [
        (re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'), '[E-MAIL REDACTED]'),
        (re.compile(r'\b[A-Z]{2}\d{2}\s?\d{4}\s?\d{4}\s?\d{4}\s?\d{4}\s?\d{0,2}\b'), '[IBAN REDACTED]'),
        (re.compile(r'\b\d{2,3}/\d{3}/\d{4,5}\b'), '[STEUERNR REDACTED]'),
        (re.compile(r'\bHR[AB]\s*\d+\b'), '[HANDELSREG REDACTED]'),
        (re.compile(r'\b[A-ZÄÖÜ][a-zäöüß]+(?:straße|strasse|str\.|weg|gasse|platz|allee|damm|ring|ufer)\s*\d+\s*[a-zA-Z]?\b', re.IGNORECASE), '[ADRESSE REDACTED]'),
        (re.compile(r'\b\d{2}\s?\d{6}\s?[A-Z]\s?\d{3}\b'), '[SOZVERSNR REDACTED]'),
    ]
    if sensitivity in ("standard", "aggressiv"):
        patterns.append(
            (re.compile(r'\b(?:\+\d{1,3}\s?)?(?:\(0\)\s?|\d{2,5}[\s/-])\d{2,5}[\s/-]?\d{2,8}\b'), '[TEL REDACTED]')
        )
        patterns.append(
            (re.compile(r'\b\d{4,5}\s+[A-ZÄÖÜ][a-zäöüß]+(?:\s+[a-zäöüß]+)?\b'), '[PLZ-ORT REDACTED]')
        )
    patterns.append(
        (re.compile(r'(?:geb(?:oren)?\.?\s*(?:am\s*)?|Geburtsdatum\s*:?\s*|geboren\s+am\s+|\*\s*)(\d{1,2}\.\d{1,2}\.\d{2,4})', re.IGNORECASE), '[GEBURTSDATUM REDACTED]')
    )
    if sensitivity == "aggressiv":
        patterns.append(
            (re.compile(r'\b\d{1,2}\.\d{1,2}\.\d{2,4}\b'), '[DATUM REDACTED]')
        )
    return patterns


ACTIVE_REGEX_PATTERNS = get_regex_patterns("standard")


def set_sensitivity(sensitivity):
    global ACTIVE_REGEX_PATTERNS
    ACTIVE_REGEX_PATTERNS = get_regex_patterns(sensitivity)


def redact_regex(text):
    # Grundbuch-Brüche schützen (z.B. 128/542, 1/3) — temporär ersetzen
    fraction_pattern = re.compile(r'\b(\d{1,6}/\d{1,6})\b')
    fractions = {}
    for i, match in enumerate(fraction_pattern.finditer(text)):
        placeholder = f"__FRACTION_{i}__"
        fractions[placeholder] = match.group()
    for placeholder, original in fractions.items():
        text = text.replace(original, placeholder, 1)

    # Regex-Schwärzung anwenden
    for pattern, replacement in ACTIVE_REGEX_PATTERNS:
        text = pattern.sub(replacement, text)

    # Brüche wiederherstellen
    for placeholder, original in fractions.items():
        text = text.replace(placeholder, original)

    return text


# ==================== NER-ERKENNUNG ====================

# Mapping von Flair-Legal-Tags zu unseren Standard-Labels
FLAIR_LEGAL_TAG_MAP = {
    "PER": "PER",
    "ORG": "ORG",
    "LOC": "LOC",
    # flair/ner-german-legal kann auch spezifischere Tags haben:
    "AN": "PER",   # Anwalt
    "UN": "ORG",   # Unternehmen
    "RS": None,     # Rechtssache/Gesetzesreferenz — NICHT schwärzen
    "GS": None,     # Gesetz — NICHT schwärzen
    "LD": "LOC",   # Land
    "ST": "LOC",   # Stadt
    "STR": "LOC",  # Straße
    "LDS": "LOC",  # Landschaft
    "RR": None,     # Rechtsreferenz — NICHT schwärzen
    "INN": None,    # Institution — prüfen per Whitelist
    "LIT": None,    # Literatur — NICHT schwärzen
    "MRK": "ORG",  # Marke
    "EUN": "ORG",  # EU-Norm — prüfen
}

# Standard-Tags von flair/ner-german-large
FLAIR_STANDARD_TAG_MAP = {
    "PER": "PER",
    "ORG": "ORG",
    "LOC": "LOC",
    "MISC": None,  # Sonstiges — nicht automatisch schwärzen
}


def _extract_entities_flair(text, mapper):
    """Extrahiert Entities mit Flair (legal + large Modell kombiniert)."""
    from flair.data import Sentence

    entities = []
    seen_spans = set()

    # 1. Legal-Modell (spezialisiert auf Rechtstexte)
    if _flair_tagger_legal:
        sentence_legal = Sentence(text)
        _flair_tagger_legal.predict(sentence_legal)

        for entity in sentence_legal.get_spans("ner"):
            tag = entity.get_label("ner").value
            score = entity.get_label("ner").score
            mapped_label = FLAIR_LEGAL_TAG_MAP.get(tag)

            if mapped_label is None:
                continue  # Gesetzesreferenzen etc. überspringen

            span_key = (entity.start_position, entity.end_position)
            if span_key not in seen_spans:
                seen_spans.add(span_key)
                entities.append({
                    "start": entity.start_position,
                    "end": entity.end_position,
                    "text": entity.text,
                    "label": mapped_label,
                    "score": score,
                    "source": "legal"
                })

    # 2. Large-Modell (allgemein stark, ergänzend)
    if _flair_tagger_large:
        sentence_large = Sentence(text)
        _flair_tagger_large.predict(sentence_large)

        for entity in sentence_large.get_spans("ner"):
            tag = entity.get_label("ner").value
            score = entity.get_label("ner").score
            mapped_label = FLAIR_STANDARD_TAG_MAP.get(tag)

            if mapped_label is None:
                continue

            span_key = (entity.start_position, entity.end_position)
            if span_key not in seen_spans:
                seen_spans.add(span_key)
                entities.append({
                    "start": entity.start_position,
                    "end": entity.end_position,
                    "text": entity.text,
                    "label": mapped_label,
                    "score": score,
                    "source": "large"
                })

    return entities


def _extract_entities_spacy(text, mapper):
    """Extrahiert Entities mit spaCy (schnellere Alternative)."""
    doc = _spacy_nlp(text)
    entities = []
    for ent in doc.ents:
        if ent.label_ in ("PER", "ORG", "LOC"):
            # Heuristischer Confidence-Score für spaCy
            score = 0.85
            word_count = len(ent.text.split())
            if word_count >= 2:
                score += 0.05
            if ent.label_ == "PER" and all(w[0].isupper() for w in ent.text.split() if w):
                score += 0.05
            if len(ent.text.strip()) <= 2:
                score -= 0.30

            entities.append({
                "start": ent.start_char,
                "end": ent.end_char,
                "text": ent.text,
                "label": ent.label_,
                "score": score,
                "source": "spacy"
            })
    return entities


def extract_entities(text, mapper):
    """Extrahiert Entities mit der aktiven NER-Engine."""
    if _nlp_engine == "flair":
        return _extract_entities_flair(text, mapper)
    else:
        return _extract_entities_spacy(text, mapper)


def redact_ner(text, mapper):
    """
    Erkennt PER, ORG und LOC-Entities und ersetzt sie mit konsistenten Platzhaltern.
    Berücksichtigt Whitelist, Confidence-Threshold und False-Positive-Heuristiken.
    """
    if not text or not text.strip():
        return text

    entities = extract_entities(text, mapper)

    # Nach Position sortieren (rückwärts) für sichere Ersetzung
    entities.sort(key=lambda x: x["start"], reverse=True)

    redacted = text
    for ent in entities:
        ent_text = ent["text"]
        ent_label = ent["label"]
        score = ent["score"]

        # 0. Gelernt: NIE schwärzen
        if is_learned_never_redact(ent_text):
            mapper.skipped_whitelist.append((ent_text, ent_label))
            continue

        # 1. Whitelist
        if is_whitelisted(ent_text, ent_label):
            mapper.skipped_whitelist.append((ent_text, ent_label))
            continue

        # 2. Juristische Personen bei konservativ nicht schwärzen
        if ent_label == "ORG" and mapper.sensitivity == "konservativ":
            mapper.skipped_org_juristic.append((ent_text, ent_label))
            continue

        # 3. Confidence
        if score < mapper.confidence_threshold:
            mapper.skipped_low_confidence.append((ent_text, ent_label, score))
            continue

        # 4. False-Positive-Heuristik
        if _should_skip_entity(ent_text, ent_label):
            continue

        placeholder = mapper.get_placeholder(ent_text, ent_label)
        if placeholder:
            redacted = redacted[:ent["start"]] + placeholder + redacted[ent["end"]:]

    return redacted


def _apply_always_redact(text, mapper):
    """Wendet die 'immer schwärzen'-Liste an — unabhängig von NER."""
    always = get_learned_always_redact()
    for label, terms in always.items():
        for term in terms:
            if term in text:
                placeholder = mapper.get_placeholder(term, label)
                if placeholder:
                    text = text.replace(term, placeholder)
    return text


def redact_text_full(text, mapper):
    """Wendet zuerst Regex, dann NER, dann gelernte 'immer schwärzen' an."""
    text = redact_regex(text)
    text = redact_ner(text, mapper)
    text = _apply_always_redact(text, mapper)
    return text


# ==================== DOCX-VERARBEITUNG ====================

def redact_paragraph(para, mapper):
    full_text = para.text
    if not full_text or not full_text.strip():
        return

    redacted_full = redact_text_full(full_text, mapper)
    if redacted_full == full_text:
        return

    runs = para.runs
    if len(runs) == 0:
        return
    if len(runs) == 1:
        runs[0].text = redacted_full
        return

    for r in runs:
        if r.text and r.text.strip():
            r.text = redact_text_full(r.text, mapper)


def process_tables(doc, mapper):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    redact_paragraph(para, mapper)


def process_headers_and_footers(doc, mapper):
    for section in doc.sections:
        for para in section.header.paragraphs:
            redact_paragraph(para, mapper)
        for para in section.footer.paragraphs:
            redact_paragraph(para, mapper)


def process_footnotes(doc, mapper):
    if hasattr(doc, "footnotes"):
        try:
            for footnote in doc.footnotes.part.document.paragraphs:
                redact_paragraph(footnote, mapper)
        except Exception:
            pass


def process_docx(file_path, output_path, mapper=None):
    if mapper is None:
        mapper = EntityMapper()

    doc = Document(file_path)

    for para in doc.paragraphs:
        redact_paragraph(para, mapper)

    process_tables(doc, mapper)
    process_headers_and_footers(doc, mapper)
    process_footnotes(doc, mapper)

    doc.save(output_path)
    print(f"DOCX erfolgreich geschwärzt: {output_path}")
    return mapper


def process_docx_api(file_path, output_path):
    doc = Document(file_path)

    for para in doc.paragraphs:
        if para.text and para.text.strip():
            redacted = redact_text_api(para.text)
            if len(para.runs) == 1:
                para.runs[0].text = redacted
            elif len(para.runs) > 1:
                para.runs[0].text = redacted
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.text = redacted

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        redacted = redact_text_api(para.text)
                        if len(para.runs) >= 1:
                            para.runs[0].text = redacted
                            for r in para.runs[1:]:
                                r.text = ""

    doc.save(output_path)
    print(f"API-basierte Redaktion abgeschlossen: {output_path}")


if __name__ == '__main__':
    input_file = 'input.docx'
    output_file = 'output_redacted.docx'
    process_docx(input_file, output_file)
